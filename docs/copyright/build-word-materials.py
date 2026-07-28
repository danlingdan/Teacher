from __future__ import annotations

import os
import subprocess
import textwrap
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[2]
OUTPUT_DIR = ROOT / "output" / "copyright-word"
QA_DIR = ROOT / "target" / "copyright-word-qa"

SOFTWARE_NAME = "SQLTeacher教学与智能练习软件"
VERSION = "V1.1.0"
BASELINE = "163aa1b15d6c2e6455f56908cdfb55a31e9d5d86"
COPYRIGHT_OWNER = "河南科技大学"

MANUAL_PATH = OUTPUT_DIR / f"{SOFTWARE_NAME}用户手册{VERSION}.docx"
SOURCE_PATH = OUTPUT_DIR / f"{SOFTWARE_NAME}源代码{VERSION}.docx"

BLUE = "1F4E79"
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "6B7280"
RED = "9B1C1C"


@dataclass(frozen=True)
class ManualPage:
    title: str
    lead: str
    steps: tuple[str, ...] = ()
    screenshot: str | None = None
    screenshot_hint: str | None = None
    note: str | None = None


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_run_font(run, name="Microsoft YaHei", size=10.5, bold=False, color="000000") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def configure_a4(section, top=1.9, bottom=1.7, left=2.2, right=2.0) -> None:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)


def add_page_field(paragraph, total: int | None = None) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    set_run_font(run, size=8.5, color=MID_GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    run = paragraph.add_run(f" 页" + (f" / 共 {total} 页" if total else ""))
    set_run_font(run, size=8.5, color=MID_GRAY)


def configure_manual_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, before, after, color in (
        ("Title", 26, 0, 10, BLUE),
        ("Subtitle", 14, 0, 8, MID_GRAY),
        ("Heading 1", 16, 0, 10, BLUE),
        ("Heading 2", 13, 10, 6, BLUE),
        ("Heading 3", 11.5, 8, 4, "1F3A5F"),
    ):
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = style_name != "Subtitle"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_manual_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run(f"{SOFTWARE_NAME}  {VERSION}  |  用户手册")
    set_run_font(run, size=8.5, color=MID_GRAY)
    p_pr = header._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), "D5DCE5")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    add_page_field(section.footer.paragraphs[0], 25)


def find_font(size: int):
    for candidate in (
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
    ):
        if Path(candidate).exists():
            try:
                return ImageFont.truetype(candidate, size=size, index=0)
            except OSError:
                continue
    return ImageFont.load_default()


def make_placeholder(number: int, title: str, hint: str) -> Path:
    QA_DIR.mkdir(parents=True, exist_ok=True)
    path = QA_DIR / f"placeholder-{number:02d}.png"
    width, height = 1600, 820
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    for offset in range(0, 8):
        draw.rounded_rectangle(
            (18 + offset, 18 + offset, width - 18 - offset, height - 18 - offset),
            radius=22,
            outline="#8FA7BF" if offset % 2 == 0 else "white",
            width=2,
        )
    title_font = find_font(48)
    body_font = find_font(30)
    small_font = find_font(24)
    draw.text((width / 2, 210), f"截图占位 {number:02d}", font=title_font, fill="#1F4E79", anchor="mm")
    draw.text((width / 2, 305), title, font=body_font, fill="#27364A", anchor="mm")
    wrapped = textwrap.wrap(hint, width=38)
    y = 405
    for line in wrapped[:3]:
        draw.text((width / 2, y), line, font=small_font, fill="#64748B", anchor="mm")
        y += 48
    draw.text((width / 2, 670), "在 Word 中选中本图，使用“更改图片”替换为完整、脱敏的运行界面", font=small_font, fill="#9B1C1C", anchor="mm")
    image.save(path, quality=95)
    return path


def add_callout(doc: Document, label: str, text: str, fill=LIGHT_BLUE, color=BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"{label}：")
    set_run_font(run, size=9.5, bold=True, color=color)
    run = p.add_run(text)
    set_run_font(run, size=9.5, color="263648")


def create_decimal_numbering(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    level.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    level.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    override.append(start_override)
    num.append(override)
    numbering.append(num)
    return num_id


def add_steps(doc: Document, steps: tuple[str, ...]) -> None:
    sequence_counter = getattr(doc, "_sqlteacher_step_sequence", 0) + 1
    setattr(doc, "_sqlteacher_step_sequence", sequence_counter)
    sequence_name = f"SQLTeacherStep{sequence_counter}"
    for step in steps:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.45)
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), f"SEQ {sequence_name} \\* ARABIC")
        p._p.append(fld)
        marker = p.add_run(". ")
        set_run_font(marker, size=10.2)
        run = p.add_run(step)
        set_run_font(run, size=10.2)


def add_manual_page(doc: Document, page: ManualPage, index: int) -> None:
    heading = doc.add_paragraph(style="Heading 1")
    heading.paragraph_format.keep_with_next = True
    run = heading.add_run(page.title)
    set_run_font(run, size=16, bold=True, color=BLUE)

    lead = doc.add_paragraph()
    lead.paragraph_format.first_line_indent = Cm(0.74)
    lead.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = lead.add_run(page.lead)
    set_run_font(run, size=10.5)

    if page.steps:
        add_steps(doc, page.steps)

    if page.screenshot and page.screenshot_hint:
        image_path = make_placeholder(index, page.screenshot, page.screenshot_hint)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(3)
        p.add_run().add_picture(str(image_path), width=Cm(14.8), height=Cm(7.6))
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_after = Pt(5)
        run = caption.add_run(f"图 {index - 4}  {page.screenshot}（截图待补）")
        set_run_font(run, size=8.8, color=MID_GRAY)

    if page.note:
        add_callout(doc, "说明", page.note)


def manual_pages() -> list[ManualPage]:
    return [
        ManualPage(
            "1 软件概述与操作流程",
            "SQLTeacher面向数据库课程教学、实验练习和课后辅导，将本地SQLite教学数据库、安全SQL执行、闯关练习、确定性评测、学习记录、课程知识检索和AI辅助整合在一个桌面应用中。核心功能本地优先，网络AI和云端教学均为可选增强。",
            ("启动软件并选择访客或账号身份。", "从主导航进入目标功能页面。", "输入SQL或自然语言需求，先完成本地安全校验。", "查看结果、评测反馈或学习记录；需要时再执行同步与导出。"),
            note="所有AI生成SQL仅作为草案，不能直接访问数据库连接或自动执行。",
        ),
        ManualPage(
            "2 安装、启动与首次初始化",
            "正式Windows安装包内含Java运行环境，普通用户无需另行安装JDK。首次启动时软件会创建本地应用数据库和教学演示数据库；开发人员也可在JDK 21与Maven环境中运行。建议显示分辨率不低于1366×768。",
            ("运行SQLTeacher安装程序并按向导完成安装。", "从开始菜单或桌面快捷方式启动软件。", "等待首次初始化完成；如提示失败，检查磁盘空间和用户目录写入权限。"),
            "软件启动或登录入口完整界面",
            "应包含软件名称、窗口边框和完整导航区域，不得只截局部。",
            "Windows正式版用户数据默认保存在%LOCALAPPDATA%\\SQLTeacher，卸载时默认保留用户数据。",
        ),
        ManualPage(
            "3 登录、注册与访客进入",
            "软件启动后显示身份门禁。用户可使用邮箱和密码登录，学生可按界面提示注册账号，也可选择访客进入。访客记录只保留在本机，不上传到账号；登录后页面会显示当前身份并按角色提供相应入口。",
            ("输入邮箱与密码后选择登录。", "首次使用的学生可选择注册并按提示完成信息。", "无需云端功能时选择访客进入。", "需要切换账号时，从主窗口右上角选择切换身份。"),
            "登录、注册与访客入口",
            "不要填写真实密码；如显示邮箱，应使用测试账号或完成脱敏。",
            "教师和管理员能力同时受客户端页面与服务端权限校验约束，隐藏按钮不是授权边界。",
        ),
        ManualPage(
            "4 首页与主导航",
            "主窗口左侧提供首页、自由SQL、闯关练习、教师题库、教师看板、课程知识、AI助手、表结构、设置与数据、云端教学等入口。中央区域显示当前页面，右上角显示身份；耗时操作期间会出现加载状态，避免重复提交。",
            ("在左侧导航选择功能。", "查看顶部身份标签，确认当前为访客、学生、教师或管理员。", "页面加载期间等待遮罩消失，再继续下一步操作。"),
            "首页与完整主导航",
            "截图应完整保留左侧功能栏、软件名称、当前身份和中央内容区域。",
        ),
        ManualPage(
            "5 自由SQL页面与示例填充",
            "自由SQL页面用于输入、检查和执行SQL。页面提供内置示例按钮，点击只会把示例语句填入编辑区，不会立即执行。用户可在提交前修改表名、字段、筛选条件和排序方式。",
            ("进入“自由 SQL”。", "选择示例SQL或在编辑区手工输入语句。", "核对当前数据库、语句范围和预期结果。", "点击执行，进入风险分析与结果展示流程。"),
            "自由SQL初始页面",
            "应展示示例区、SQL编辑区、执行按钮及结果区域的完整布局。",
            "示例建议使用内置student表，避免展示外部数据库地址、账号或真实业务数据。",
        ),
        ManualPage(
            "6 查询执行与结果阅读",
            "只读查询通过安全检查后由数据库适配器执行。结果区域显示列名、数据行、执行耗时和截断提示；无结果时保持列结构并给出说明；语法或数据库错误会转换为学习者可理解的提示，不直接暴露内部堆栈。",
            ("执行SELECT查询。", "按表头核对返回列。", "检查结果行、排序和筛选条件。", "如出现截断提示，增加更精确的WHERE条件，不要依赖无限制返回。"),
            "查询成功结果",
            "使用演示数据展示列名、若干结果行、耗时和状态提示。",
        ),
        ManualPage(
            "7 SQL风险识别与高风险确认",
            "软件在执行前识别查询、数据修改、结构修改、多语句和禁止语句。只读查询通常允许执行；UPDATE、DELETE、ALTER、DROP等高风险语句需要用户明确确认；多语句以及DROP DATABASE、GRANT、REVOKE等默认拦截。",
            ("提交SQL后阅读风险等级和原因。", "高风险操作先核对目标表、条件与影响范围。", "确认无误后选择确认执行；不确定时选择取消。", "被禁止的语句需要改写，不能通过拼接或AI绕过。"),
            "高风险SQL确认对话框",
            "对话框应显示完整待执行SQL、风险说明、取消与确认按钮。",
            "确认按钮仅表示用户确认当前操作，并不替代数据库权限、备份或审计。",
        ),
        ManualPage(
            "8 禁止语句、错误与安全拦截",
            "当SQL包含多条语句、禁止级操作或不符合安全规则时，软件在访问数据库前终止流程并显示原因。用户应根据提示拆分、改写或放弃操作。AI草案与手写SQL使用同一套Java侧风险规则。",
            ("阅读拦截原因和命中的规则。", "删除额外分号或拆分多语句。", "对禁止级操作使用数据库管理员批准的独立工具和流程，不在教学软件中绕过。", "修改后重新提交并再次检查。"),
            "禁止SQL拦截提示",
            "可使用多语句或DROP DATABASE示例，但不得连接真实生产数据库。",
        ),
        ManualPage(
            "9 表结构浏览",
            "表结构页面展示当前数据库中的表、字段、数据类型、可空性和主键标记。学习者可先查看元数据再编写SQL；AI助手也只使用必要的受控结构信息，不读取或执行任意数据库操作。",
            ("进入“表结构”。", "选择目标表。", "查看字段名称、类型、主键和可空性。", "回到自由SQL或AI助手，按实际结构编写语句。"),
            "表结构浏览页面",
            "应展示完整表列表和字段明细，不包含外部数据库的敏感命名。",
        ),
        ManualPage(
            "10 数据库连接设置与测试",
            "设置与数据中的连接管理支持SQLite、MySQL和MariaDB的受控配置。连接档案保存非敏感参数；密码只用于本次连接测试，不长期保存。外部数据库应使用最小权限账号，且用户需独立复核所有修改操作。",
            ("新增连接并选择数据库类型。", "填写路径或主机、端口、数据库名和用户名。", "按需输入临时密码并选择“测试连接”。", "测试成功后保存非敏感配置并设为当前连接。"),
            "数据库连接设置",
            "主机、用户名、路径和测试密码必须使用虚构值或遮盖；保留完整窗口。",
            "网络数据库不可达时，应分别检查网络、端口、TLS和账号权限，不把密码写入截图或日志。",
        ),
        ManualPage(
            "11 闯关练习目录与题目选择",
            "学生在闯关练习中查看已启用题目、知识点、难度和完成状态。选择题目后可阅读任务说明、目标数据集和限制，再进入隔离练习会话。题目是否可见由题库状态和当前身份共同决定。",
            ("进入“闯关练习”。", "按知识点或难度浏览题目。", "选择题目并阅读说明。", "开始练习后在隔离数据环境中编写SQL。"),
            "闯关练习目录",
            "展示题目列表、难度、知识点及状态；不要显示真实学生信息。",
        ),
        ManualPage(
            "12 SQL练习、分级提示与提交评测",
            "练习页面提供SQL编辑、分级提示和提交评测。系统按照题目定义的列、数据行、行序、指定行数和关键结构等规则进行确定性比较，返回通过状态、差异信息或改进建议，不由AI决定最终判题。",
            ("阅读题目并独立编写SQL。", "遇到困难时逐级查看提示。", "执行或提交前检查SQL风险。", "提交评测并根据差异反馈修改，直至通过。"),
            "学生练习与评测结果",
            "同一画面应包含题目、SQL编辑区、提示入口和评测反馈。",
            "练习数据集与自由SQL当前连接相互隔离，避免题目之间或用户操作之间串扰。",
        ),
        ManualPage(
            "13 教师题库管理",
            "教师题库用于创建、编辑、启用或停用练习题，维护知识点、难度、数据集、参考SQL、评测规则与分级提示。题包支持导入和导出；导入发生冲突时整体回滚，避免留下半完成数据。",
            ("以教师身份进入“教师题库”。", "新增或选择题目并填写定义。", "配置参考SQL、评测规则和提示。", "保存后启用题目；需要迁移时使用题包导入或导出。"),
            "教师题库编辑页面",
            "展示题目定义和保存状态；参考SQL只使用教学示例数据。",
        ),
        ManualPage(
            "14 教师看板、筛选与CSV导出",
            "教师看板汇总本地学习事件和练习进度，可按日期、题目、知识点与错误类型组合筛选，并导出UTF-8 CSV。教师可据此识别常见困难，但导出文件仅应用于授权的课程教学和分析。",
            ("进入“教师看板”。", "设置日期、题目、知识点或错误类型筛选条件。", "查看统计与明细。", "确认范围后导出CSV，并按学校要求妥善保存。"),
            "教师看板筛选结果",
            "学生姓名、学号、邮箱和真实学习记录必须使用测试数据或脱敏。",
        ),
        ManualPage(
            "15 课程知识导入与全文检索",
            "课程知识页面可导入经授权的本地课程文档，建立SQLite FTS5全文索引，并按关键词检索内容和展示来源。教师可移除不再需要的索引。课程资料及检索结果应按学校的数据和版权要求管理。",
            ("选择本地课程文档并执行导入。", "等待索引完成并查看状态。", "输入关键词检索。", "核对命中片段及来源；不再使用时删除索引。"),
            "课程知识检索结果",
            "建议使用自编测试文档，展示关键词、命中片段和来源文件名。",
        ),
        ManualPage(
            "16 AI助手、模型选择与SQL草案",
            "AI助手可检测本地Ollama模型并选择可用模型，也可使用用户自行配置的兼容OpenAI Chat Completions网络服务。用户输入自然语言需求后，系统请求结构化输出，并展示SQL草案与教学解释。",
            ("进入“AI助手”并确认服务状态。", "选择可用本地模型或已配置的网络AI。", "输入清晰、单一的查询需求。", "生成后核对SQL草案、解释、表名和字段名。"),
            "AI助手生成SQL草案",
            "画面应包含模型状态、自然语言输入、SQL草案、解释和安全状态；不得出现API Key。",
        ),
        ManualPage(
            "17 AI输出校验、预览与安全降级",
            "模型输出被视为不可信输入。Java侧依次进行JSON解析、字段完整性、方言与SQL风险校验；通过后也只作为草案展示。用户须将草案带入自由SQL页面，再次检查并主动执行。服务不可用、超时或格式错误时，软件给出降级提示。",
            ("阅读草案和安全检查结果。", "发现表名、字段或条件不正确时先修改。", "将确认后的草案复制或带入自由SQL。", "AI不可用时继续手写SQL、题目练习、统计或知识检索。"),
            "AI不可用或校验失败状态",
            "可展示本地模型未启动或格式校验失败，但不得暴露内部地址、令牌或完整日志。",
            "网络AI的API Key仅驻留当前进程内存，不同步、不备份、不写入日志。",
        ),
        ManualPage(
            "18 云端账号、班级与成员管理",
            "启用云端服务后，教师可创建或管理班级成员，学生可加入班级并查看所属资源。服务端按登录身份、角色与成员关系校验权限。云端不可用时，本地SQLite练习、确定性评测和知识检索仍可继续。",
            ("登录测试账号并进入“云端教学”。", "教师创建班级或维护成员；学生按指引加入班级。", "确认当前角色与可见资源。", "退出后检查身份标签和本地功能。"),
            "云端班级与成员页面",
            "使用测试班级与虚构成员；不展示真实服务器地址、访问令牌或学生资料。",
        ),
        ManualPage(
            "19 教学任务与学习事件同步",
            "教师可发布班级任务，学生查看任务并完成本地练习。学习事件采用幂等上传和增量下载，本地按账号标记事件所有者；访客及其他账号的事件不会上传到当前账号。网络失败时保留本地记录，恢复后可重试同步。",
            ("教师填写任务标题、关联题目和可选截止时间并发布。", "学生在所属班级查看任务。", "完成练习后返回云端教学并选择立即同步。", "查看上传、下载或失败提示；网络恢复后重试。"),
            "任务发布与同步状态",
            "同一画面可展示任务列表和同步结果；时间、班级与账号均使用测试数据。",
        ),
        ManualPage(
            "20 版本、本地备份与数据恢复",
            "版本与本地数据页面显示软件版本、用户数据目录和备份列表。软件在升级前自动备份，也支持立即备份、SQLite完整性检查后恢复所选备份，以及一键复原内置演示数据库。恢复操作前应确认影响范围。",
            ("进入“设置与数据”中的版本与本地数据。", "选择“立即备份”并确认新备份出现。", "恢复前核对备份时间和数据范围。", "需要重置教学样例时选择恢复内置演示库。"),
            "版本与本地数据页面",
            "应显示V1.1.0、备份列表和操作按钮；用户目录中的个人用户名需要遮盖。",
            "重要课程和学习数据应另行保留独立副本，避免恢复或复原操作造成不可逆覆盖。",
        ),
        ManualPage(
            "21 常见问题与故障处理",
            "遇到问题时优先根据页面状态定位：启动失败通常与磁盘空间或写权限有关；SQL错误需检查关键字、字段、括号与逗号；AI不可用需检查Ollama模型或网络配置；云端失败不影响本地练习。",
            ("记录发生时间、当前页面和可复现步骤。", "根据页面提示检查输入、网络或本地服务。", "重试前确认操作不会重复创建任务或覆盖数据。", "反馈问题时仅提供脱敏截图和必要日志片段。"),
            note="不得向维护人员发送数据库密码、API Key、访问令牌、完整应用数据库或含个人信息的未脱敏截图。",
        ),
        ManualPage(
            "22 安全使用与提交前核对",
            "本手册描述V1.1.0冻结版本已经实现的功能。正式用于软件著作权登记前，应将全部占位图替换为同一版本、同一软件名称的完整界面截图，并逐项检查申请表、说明书和源代码材料的一致性。",
            ("确认封面、页眉和源代码页眉的软件全称完全一致。", "确认所有页面版本号均为V1.1.0，未混入后续版本界面。", "确认一级功能完整展示，关键二级操作有对应截图和文字说明。", "检查截图无密码、密钥、令牌、真实学生信息和内部服务地址。", "在Word中更新并检查页码，另存一份最终PDF进行逐页复核。"),
            note="开发完成日期、发表日期、著作权人证件和签章流程应以申请主体最终确认及受理系统当期要求为准。",
        ),
    ]


def build_manual() -> None:
    doc = Document()
    configure_a4(doc.sections[0])
    configure_manual_styles(doc)
    add_manual_header_footer(doc)
    props = doc.core_properties
    props.title = f"{SOFTWARE_NAME}用户手册{VERSION}"
    props.author = COPYRIGHT_OWNER
    props.subject = "计算机软件著作权登记操作说明书"

    # Page 1: editorial cover.
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(SOFTWARE_NAME)
    set_run_font(run, size=25, bold=True, color=BLUE)
    p.paragraph_format.space_after = Pt(12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("用户手册")
    set_run_font(run, size=20, bold=True, color="27364A")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(VERSION)
    set_run_font(run, size=15, bold=True, color=BLUE)
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"著作权人：{COPYRIGHT_OWNER}")
    set_run_font(run, size=11, color="27364A")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("文档版本：1.0  |  对应软件版本：1.1.0")
    set_run_font(run, size=9.5, color=MID_GRAY)
    doc.add_page_break()

    # Page 2: document information.
    p = doc.add_paragraph(style="Heading 1")
    p.add_run("文档信息与截图替换说明")
    rows = [
        ("软件全称", SOFTWARE_NAME),
        ("软件简称", "SQLTeacher"),
        ("版本号", VERSION),
        ("著作权人", COPYRIGHT_OWNER),
        ("适用平台", "Windows 10/11 64位"),
        ("冻结基线", f"Git提交 {BASELINE[:12]}（v1.1.0）"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_geometry(table, [2700, 6660])
    for row, (label, value) in zip(table.rows, rows):
        set_cell_shading(row.cells[0], LIGHT_GRAY)
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        r = row.cells[0].paragraphs[0].add_run(label)
        set_run_font(r, size=9.5, bold=True, color=BLUE)
        r = row.cells[1].paragraphs[0].add_run(value)
        set_run_font(r, size=9.5)
    doc.add_paragraph()
    add_callout(doc, "截图替换", "本手册已预留编号截图框。选中占位图后，在Word中使用“图片格式→更改图片”替换；保持原尺寸和图注即可。截图必须显示完整运行窗口、软件名称和功能导航，并使用测试数据或完成脱敏。", fill="FFF8E8", color="7A5A00")
    add_callout(doc, "一致性要求", "正式提交前，申请表、手册封面、页眉、源代码页眉中的软件全称与版本号必须完全一致。", fill=LIGHT_BLUE, color=BLUE)
    doc.add_page_break()

    # Page 3: static table of contents.
    p = doc.add_paragraph(style="Heading 1")
    p.add_run("目录")
    titles = [page.title for page in manual_pages()]
    for i, title in enumerate(titles, start=4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        left = p.add_run(title)
        set_run_font(left, size=9.8, color="27364A")
        dots = "." * max(8, 34 - len(title))
        mid = p.add_run(f" {dots} ")
        set_run_font(mid, name="Courier New", size=8.5, color="A0A8B3")
        right = p.add_run(str(i))
        set_run_font(right, size=9.8, bold=True, color=BLUE)
    doc.add_page_break()

    pages = manual_pages()
    for offset, page in enumerate(pages, start=4):
        add_manual_page(doc, page, offset)
        if offset < 25:
            doc.add_page_break()

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(MANUAL_PATH)


def git_text(path: str) -> str:
    result = subprocess.run(
        ["git", "show", f"{BASELINE}:{path}"],
        cwd=ROOT,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        check=True,
    )
    return result.stdout.decode("utf-8", errors="replace").replace("\r\n", "\n")


def source_display_lines() -> tuple[list[str], int, int]:
    result = subprocess.run(
        ["git", "ls-tree", "-r", "--name-only", BASELINE, "--", "src/main/java"],
        cwd=ROOT,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        check=True,
        text=True,
        encoding="utf-8",
    )
    paths = sorted(path for path in result.stdout.splitlines() if path.endswith(".java"))
    raw_lines: list[str] = []
    display_lines: list[str] = []
    for path in paths:
        marker = f"// ===== 文件：{path} ====="
        raw_lines.append(marker)
        display_lines.append(marker)
        for raw in git_text(path).splitlines():
            raw_lines.append(raw)
            if not raw:
                display_lines.append(" ")
                continue
            wrapped = textwrap.wrap(
                raw.expandtabs(4),
                width=128,
                replace_whitespace=False,
                drop_whitespace=False,
                break_long_words=True,
                break_on_hyphens=False,
            ) or [" "]
            display_lines.append(wrapped[0])
            display_lines.extend("    > " + item for item in wrapped[1:])
    return display_lines, len(paths), len(raw_lines)


def add_source_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run(f"{SOFTWARE_NAME} {VERSION} 源代码")
    set_run_font(run, size=8, color="404040")
    p_pr = header._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "A6A6A6")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    add_page_field(section.footer.paragraphs[0], 60)


def build_source() -> tuple[int, int, int]:
    lines, file_count, raw_line_count = source_display_lines()
    if len(lines) < 3000:
        selected = lines
    else:
        selected = lines[:1500] + lines[-1500:]

    doc = Document()
    configure_a4(doc.sections[0], top=1.55, bottom=1.35, left=1.35, right=1.25)
    add_source_header_footer(doc)
    props = doc.core_properties
    props.title = f"{SOFTWARE_NAME}源代码{VERSION}"
    props.author = COPYRIGHT_OWNER
    props.subject = "计算机软件著作权登记源程序鉴别材料"

    style = doc.styles["Normal"]
    style.font.name = "Courier New"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(5.8)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    style.paragraph_format.line_spacing = Pt(10.1)

    for page_no in range((len(selected) + 49) // 50):
        if page_no:
            doc.add_page_break()
        start = page_no * 50
        block = selected[start:start + 50]
        for row_no, line in enumerate(block, start=1):
            p = doc.add_paragraph()
            p.paragraph_format.keep_together = True
            p.paragraph_format.keep_with_next = False
            run = p.add_run(f"{row_no:02d}  {line}")
            set_run_font(run, name="Courier New", size=5.8, color="000000")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(SOURCE_PATH)
    return file_count, raw_line_count, len(lines)


def main() -> None:
    build_manual()
    file_count, raw_line_count, display_count = build_source()
    print(f"Created: {MANUAL_PATH}")
    print(f"Created: {SOURCE_PATH}")
    print(f"Baseline: {BASELINE}; Java files={file_count}; raw rows with file markers={raw_line_count}; display rows={display_count}")


if __name__ == "__main__":
    main()
