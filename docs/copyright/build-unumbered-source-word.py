from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[2]
INPUT_PATH = (
    ROOT
    / "output"
    / "copyright-plain"
    / "SQLTeacher源程序鉴别材料V1.1.0-纯文本.txt"
)
OUTPUT_PATH = (
    ROOT
    / "output"
    / "copyright-word"
    / "SQLTeacher源程序鉴别材料V1.1.0-无行号纯代码.docx"
)


def main() -> None:
    source_lines = INPUT_PATH.read_text(encoding="utf-8").splitlines()
    if len(source_lines) != 3000:
        raise RuntimeError(f"Expected 3000 source lines, found {len(source_lines)}")

    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.3)
    section.bottom_margin = Cm(1.3)
    section.left_margin = Cm(1.3)
    section.right_margin = Cm(1.3)

    style = document.styles["Normal"]
    style.font.name = "Courier New"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(6.5)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    style.paragraph_format.line_spacing = Pt(8.5)

    for line in source_lines:
        paragraph = document.add_paragraph()
        paragraph.add_run(line)

    document.core_properties.title = "SQLTeacher源程序鉴别材料V1.1.0-无行号纯代码"
    document.core_properties.subject = "无行号、无页眉页脚、无续行标记的纯源代码"
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    print(f"Created: {OUTPUT_PATH}")
    print(f"Pure source lines: {len(source_lines)}")


if __name__ == "__main__":
    main()
