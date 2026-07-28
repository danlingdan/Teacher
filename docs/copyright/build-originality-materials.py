from __future__ import annotations

import subprocess
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT_DIR = ROOT / "output" / "copyright-originality"
SCREENSHOT_DIR = ROOT / "docs" / "copyright" / "screenshots"
BASELINE = "163aa1b15d6c2e6455f56908cdfb55a31e9d5d86"
SOFTWARE_NAME = "SQLTeacher SQL教学与智能练习软件"
VERSION = "V1.1.0"

MANUAL_PATH = OUTPUT_DIR / f"{SOFTWARE_NAME}用户手册{VERSION}.docx"
SOURCE_PATH = OUTPUT_DIR / f"{SOFTWARE_NAME}源程序鉴别材料{VERSION}.docx"

BLUE = "2E74B5"
NAVY = "17365D"
INK = "202B38"
MUTED = "64748B"
PALE_BLUE = "E8EEF5"
PALE_PURPLE = "F1EEFF"
PALE_GREEN = "EAF5EF"
PALE_AMBER = "FFF4D6"

REPRESENTATIVE_FILES = [
    "src/main/java/com/sqlteacher/infrastructure/database/DefaultSqlRiskAnalysisService.java",
    "src/main/java/com/sqlteacher/infrastructure/database/SqlScriptSplitter.java",
    "src/main/java/com/sqlteacher/infrastructure/database/JdbcSqlExecutionService.java",
    "src/main/java/com/sqlteacher/infrastructure/database/DeterministicSqlExerciseEvaluationService.java",
    "src/main/java/com/sqlteacher/infrastructure/database/JdbcExercisePracticeService.java",
    "src/main/java/com/sqlteacher/infrastructure/ai/Nl2SqlServiceImpl.java",
    "src/main/java/com/sqlteacher/application/nl2sql/DefaultNl2SqlSafetyService.java",
    "src/main/java/com/sqlteacher/infrastructure/database/SqliteKnowledgeService.java",
    "src/main/java/com/sqlteacher/application/event/DefaultLearningEventService.java",
    "src/main/java/com/sqlteacher/infrastructure/database/JdbcLearningAnalyticsService.java",
    "src/main/java/com/sqlteacher/infrastructure/database/SqliteApplicationBackupService.java",
]


def set_run_font(run, name="Microsoft YaHei", size=10.5, bold=False, color=INK, italic=False):
    run.font.name = name
    fonts = run._element.get_or_add_rPr().rFonts
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), "Microsoft YaHei" if name == "Courier New" else name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def configure_a4(section, top=1.7, bottom=1.6, left=1.8, right=1.8):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)
    section.header_distance = Cm(0.75)
    section.footer_distance = Cm(0.75)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    set_run_font(run, size=8.5, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)
    run = paragraph.add_run(" 页")
    set_run_font(run, size=8.5, color=MUTED)


def add_bottom_border(paragraph, color="CBD5E1"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)


def configure_manual_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, NAVY, 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_header_footer(doc):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(header.add_run(f"{SOFTWARE_NAME}  {VERSION}  |  用户手册"), size=8.5, color=MUTED)
    add_bottom_border(header)
    add_page_field(section.footer.paragraphs[0])


def create_numbering(doc, fmt="bullet"):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), fmt)
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if fmt == "bullet" else "%1.")
    level.append(lvl_text)
    if fmt == "bullet":
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), "Arial")
        r_fonts.set(qn("w:hAnsi"), "Arial")
        r_pr.append(r_fonts)
        level.append(r_pr)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    level.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    override.append(start_override)
    num.append(override)
    numbering.append(num)
    return num_id


def add_list_item(doc, text, num_id):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.25)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.25
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "14")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), BLUE)
    borders.append(left)
    p_pr.append(borders)
    p.add_run(text)
    return p


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_geometry(table, widths, indent=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = OxmlElement("w:tcW")
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            tc_pr.append(tc_w)
            margins = OxmlElement("w:tcMar")
            for side in ("top", "bottom", "start", "end"):
                node = OxmlElement(f"w:{side}")
                node.set(qn("w:w"), "80" if side in ("top", "bottom") else "120")
                node.set(qn("w:type"), "dxa")
                margins.append(node)
            tc_pr.append(margins)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, PALE_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(header), size=9, bold=True, color=NAVY)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run_font(p.add_run(value), size=8.8, color=INK)
    set_table_geometry(table, widths)
    return table


def add_callout(doc, label, text, fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run(f"{label}："), size=9.5, bold=True, color=NAVY)
    set_run_font(p.add_run(text), size=9.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        set_run_font(p.add_run(bold_lead), bold=True, color=NAVY)
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)
    return p


def add_code(doc, code):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F6F8FA")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for index, line in enumerate(code.splitlines()):
        run = p.add_run(line)
        set_run_font(run, name="Courier New", size=8.5, color=INK)
        if index < len(code.splitlines()) - 1:
            run.add_break()


def add_screenshot(doc, filename, caption, width_cm=15.8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(str(SCREENSHOT_DIR / filename), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(5)
    set_run_font(cap.add_run(caption), size=8.5, color=MUTED, italic=True)


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_page_title(doc, number, title, lead):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    p.add_run(f"{number}  {title}")
    lead_p = doc.add_paragraph()
    lead_p.paragraph_format.space_after = Pt(9)
    set_run_font(lead_p.add_run(lead), size=11, color=MUTED)


def build_manual():
    doc = Document()
    configure_a4(doc.sections[0])
    configure_manual_styles(doc)
    configure_header_footer(doc)
    bullet_id = create_numbering(doc, "bullet")
    decimal_id = create_numbering(doc, "decimal")

    # 1 - Cover and product identity
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(45)
    set_run_font(p.add_run("数据库教学 · 安全执行 · 智能练习"), size=11, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    set_run_font(p.add_run(SOFTWARE_NAME), size=26, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    set_run_font(p.add_run(f"用户操作手册  {VERSION}"), size=15, color=BLUE)
    doc.add_heading("软件的核心特点", level=2)
    for text in [
        "学生输入的SQL先经过草稿、风险判断、确认、执行、结果解释和学习事件记录六个环节，再交给数据库处理。",
        "AI只生成结构化SQL草案。Java侧负责解析模型输出、检查方言并执行风险分析，模型不持有JDBC连接。",
        "练习结果由确定性规则评测，可检查列、行、顺序、行数和关键结构，避免用模型主观判断学生答案。",
        "手写SQL、练习评测和本地记录可离线使用。云端账号、班级、任务与同步为可选功能，网络中断时本地记录仍会保留。",
    ]:
        add_list_item(doc, text, bullet_id)
    add_table(doc, ["项目", "内容"], [
        ("软件版本", VERSION),
        ("运行形态", "Java 21 / JavaFX桌面应用，本地SQLite优先"),
        ("适用对象", "数据库课程学生、授课教师、实验管理员"),
        ("文档范围", "安装启动、SQL练习、安全控制、评测、知识与AI、教学协作、数据维护"),
    ], [2100, 7260])
    page_break(doc)

    # 2 - Overview
    add_page_title(doc, "1", "从学习任务到可追溯结果", "本章按一次练习的实际处理顺序，说明SQL从输入、执行到形成学习记录的完整过程。")
    doc.add_heading("1.1 典型学习闭环", level=2)
    decimal_id = create_numbering(doc, "decimal")
    for text in [
        "选择演示数据库或已配置连接，先浏览表和字段，确认练习上下文。",
        "输入单条SQL或从AI助手取得草案；系统拆分语句并识别只读查询、数据修改、结构变更、禁止操作或多语句拼接。",
        "安全规则决定直接执行、要求二次确认或阻止；通过后才进入JDBC执行路径。",
        "查询结果显示列、行、耗时和截断状态；数据库错误被转换为学生可理解的反馈。",
        "在闯关练习中提交答案，由确定性规则给出通过、差异和分级提示，并记录学习事件。",
        "教师在学情看板按题目、知识点和错误类型汇总，形成下一轮教学干预依据。",
    ]:
        add_list_item(doc, text, decimal_id)
    doc.add_heading("1.2 身份与边界", level=2)
    add_table(doc, ["身份", "主要入口", "数据边界"], [
        ("访客", "首页、自由SQL、闯关练习、AI助手、表结构", "仅本机体验，不上传账号"),
        ("学生", "访客功能、课程知识、云端班级与任务", "仅同步本人学习事件"),
        ("教师", "题库、看板、连接设置、班级成员与任务", "仅管理参与班级"),
        ("管理员", "全部功能与云端管理", "仍受服务端鉴权与审计约束"),
    ], [1450, 4300, 3610])
    add_callout(doc, "安全原则", "界面是否显示某个入口不是唯一权限控制。桌面控制器和云端服务都会再次校验身份、角色与班级成员关系。", PALE_AMBER)
    page_break(doc)

    # 3 - Login
    add_page_title(doc, "2", "登录、注册与访客模式", "启动页把账号教学与本地体验明确分开，避免访客数据误同步。")
    add_screenshot(doc, "01-login.png", "图1  V1.1.0登录页：登录、注册学生账号与访客进入共用同一入口", 14.2)
    doc.add_heading("2.1 操作步骤", level=2)
    decimal_id = create_numbering(doc, "decimal")
    for text in [
        "已有账号：填写邮箱和密码，点击“登录并进入”。",
        "新学生：填写邮箱、昵称和至少12位密码，点击“注册学生账号”。",
        "仅体验本地功能：点击“访客进入”，访客数据不会上传到云端账号。",
        "切换身份：在主界面右上角退出当前身份；服务端令牌撤销后返回登录页。",
    ]:
        add_list_item(doc, text, decimal_id)
    add_callout(doc, "隐私提示", "不要在截图或问题反馈中展示真实邮箱、密码、访问令牌、API Key或学生身份信息。", PALE_AMBER)
    page_break(doc)

    # 4 - Home
    add_page_title(doc, "3", "首页与功能导航", "首页按当前登录身份展示可用功能入口，访客、学生、教师和管理员看到的入口集合不同。")
    add_screenshot(doc, "02-home.png", "图2  V1.1.0首页：自由SQL、AI助手和表结构构成学生的基础工作区", 15.5)
    doc.add_heading("3.1 导航方式", level=2)
    add_body(doc, "顶部导航用于切换首页、自由SQL、闯关练习、AI助手、表结构及身份相关页面。窄屏时导航区域可横向滚动，耗时操作期间相关按钮会暂时禁用，防止重复提交。")
    doc.add_heading("3.2 建议的首次使用顺序", level=2)
    decimal_id = create_numbering(doc, "decimal")
    for text in [
        "先进入表结构，认识student等演示表的字段和主键。",
        "进入自由SQL，执行只读查询并观察结果列、行数和耗时。",
        "进入闯关练习，用题目要求和分级提示完成一次提交。",
        "需要时再使用AI助手生成草案，并返回自由SQL页面复核执行。",
    ]:
        add_list_item(doc, text, decimal_id)
    page_break(doc)

    # 5 - SQL practice
    add_page_title(doc, "4", "自由SQL", "所有SQL都先经过Java侧拆分与风险判断，数据库执行位于安全检查之后。")
    add_screenshot(doc, "03-sql-practice.png", "图3  V1.1.0自由SQL页：示例、编辑区、执行动作和查询结果位于同一学习上下文", 15.5)
    doc.add_heading("4.1 完成一次查询", level=2)
    decimal_id = create_numbering(doc, "decimal")
    for text in [
        "点击示例SQL或在编辑区输入一条语句。",
        "点击“执行SQL”；系统拒绝空输入和默认禁止的多语句。",
        "风险分析通过后执行查询；结果区显示列名、数据行、执行耗时和是否截断。",
        "若出现错误，根据中文提示检查表名、列名、引号、括号和筛选条件。",
    ]:
        add_list_item(doc, text, decimal_id)
    add_code(doc, "SELECT name, score\nFROM student\nWHERE score >= 60\nORDER BY score DESC;")
    page_break(doc)

    # 6 - Risk
    add_page_title(doc, "5", "SQL风险识别与执行控制", "风险识别同时处理语句边界、语句类型、禁止规则和高风险确认。例如，字符串字面量中的分号不会被当作切分点。")
    doc.add_heading("5.1 风险决策", level=2)
    add_table(doc, ["类别", "示例", "系统处理"], [
        ("只读查询", "SELECT", "通过检查后执行，并限制结果数量"),
        ("数据修改", "UPDATE / DELETE", "展示影响说明，要求明确确认"),
        ("结构变更", "ALTER / DROP TABLE", "按高风险处理，要求确认"),
        ("禁止操作", "DROP DATABASE / GRANT / REVOKE", "直接阻止，不进入JDBC"),
        ("多语句", "以分号拼接两条语句", "默认阻止，避免绕过单语句判断"),
    ], [1700, 2800, 4860])
    doc.add_heading("5.2 为什么能识别复杂输入", level=2)
    add_body(doc, "SQL拆分器按有限状态机实现。扫描过程中持续维护单引号、双引号、行注释和块注释的上下文状态；只有状态机处于语句文本区时，分号才被识别为语句边界。例如，SELECT ';'或/* 注释; */ SELECT 1中的分号不会触发切分。")
    add_body(doc, "风险分析器随后规范化SQL、识别首个有效关键字，并结合禁止模式与高风险模式返回风险等级、提示和是否允许执行。与直接按分号切分字符串的做法相比，状态机不会把字面量或注释中的分号误判为语句边界。")
    add_callout(doc, "不可绕过", "自由输入、AI草案和练习草稿最终都进入同一安全路径；用户确认高风险操作只表示理解当前影响，不会改变数据库权限。", PALE_AMBER)
    page_break(doc)

    # 7 - Exercises
    add_page_title(doc, "6", "闯关练习与确定性评测", "练习服务为每次作答创建隔离会话，判定依据来自题目规则而不是AI主观结论。")
    doc.add_heading("6.1 学生操作", level=2)
    decimal_id = create_numbering(doc, "decimal")
    for text in [
        "从目录选择题目，阅读知识点、难度、目标和可用数据集。",
        "开始会话后运行草稿；草稿仍需经过SQL安全分析。",
        "遇到困难时按层级使用提示，系统记录提示级别但不泄露密码或连接信息。",
        "提交答案后查看每个判定项的通过状态、差异信息和改进建议。",
        "重置会话只重建当前隔离练习库，不修改题库、应用库或其他会话。",
    ]:
        add_list_item(doc, text, decimal_id)
    doc.add_heading("6.2 评测维度", level=2)
    add_table(doc, ["维度", "判定内容"], [
        ("结果列", "列名、列数及必要的顺序"),
        ("结果数据", "行集合、重复行和是否要求固定顺序"),
        ("结果规模", "指定行数或最大行数"),
        ("SQL结构", "题目要求的关键结构和禁止结构"),
        ("安全状态", "多语句、禁止操作和未经确认的高风险语句"),
    ], [2200, 7160])
    add_body(doc, "评测服务执行学生SQL和参考SQL，再按题目配置比较结果列、行集合、顺序、行数及关键结构，不按SQL文本逐字匹配。大小写和多余空格等书写差异不会单独导致失败；列名是否必须一致由具体题目的结果列规则决定，因此可支持题目允许范围内的语义等价解法。")
    add_callout(doc, "原创教学机制", "系统把运行结果、评测差异、提示使用和错误类型统一记录为学习事件，使学生反馈与教师分析共享同一事实来源。", PALE_GREEN)
    page_break(doc)

    # 8 - AI
    add_page_title(doc, "7", "AI助手", "模型输出被视为不可信输入，必须经过结构化解析和本地安全检查。")
    doc.add_heading("7.1 生成流程", level=2)
    decimal_id = create_numbering(doc, "decimal")
    for text in [
        "用户描述查询目标，软件只提供完成任务所需的表结构上下文。",
        "本地Ollama或兼容OpenAI Chat Completions接口的HTTPS服务返回结构化内容。",
        "Java侧解析JSON并校验SQL、解释、方言和必要字段；格式错误进入降级状态。",
        "安全服务重新分析SQL，禁止语句或多语句不能作为可执行草案返回。",
        "界面展示SQL草案与教学解释；用户复制到自由SQL页后仍要再次检查和确认。",
    ]:
        add_list_item(doc, text, decimal_id)
    add_body(doc, "AI在本软件中只承担草案生成角色。它不接触数据库连接，输出内容与人工输入接受相同的Java侧安全检查，因此不能绕过SQL风险规则直接执行语句。")
    doc.add_heading("7.2 模型不可用时", level=2)
    add_body(doc, "当模型未安装、服务不可达、请求超时或响应格式不合格时，AI助手显示可理解的错误并停止生成。手写SQL、练习评测、课程知识、学情统计和本地备份仍可继续使用。")
    add_callout(doc, "密钥边界", "网络AI的API Key仅驻留当前进程内存，不同步、不备份、不写入日志；退出应用后需要重新提供。", PALE_AMBER)
    page_break(doc)

    # 9 - Knowledge and analytics
    add_page_title(doc, "8", "课程知识与学情分析", "本章分两部分：8.1节说明课程资料的导入与检索，8.2节和8.3节说明学习事件的记录与统计方式。")
    doc.add_heading("8.1 课程知识检索", level=2)
    add_body(doc, "教师可导入有权使用的UTF-8文本或Markdown资料。系统校验路径、文件类型、大小和编码，将内容分块写入SQLite FTS5索引；检索结果显示标题、来源和命中片段。删除文档时关联索引一并删除。")
    doc.add_heading("8.2 学习事件", level=2)
    add_table(doc, ["事件", "代表含义", "可分析维度"], [
        ("SQL_EXECUTION", "运行SQL草稿", "连接、耗时、成功、结果行数"),
        ("SQL_RISK_BLOCKED", "语句被安全规则阻止", "风险级别、语句类型"),
        ("EXERCISE_ATTEMPT", "提交练习答案", "题目、结果、耗时、错误码"),
        ("EXERCISE_HINT_USED", "使用分级提示", "题目、提示级别"),
        ("KNOWLEDGE_SEARCHED", "检索课程资料", "查询长度、命中数量"),
    ], [2350, 3150, 3860])
    add_body(doc, "每条学习事件包含事件ID、发生时间、事件类型和属性数据；登录状态下还写入所属账号，练习事件会关联会话ID。事件ID同时作为9.2节增量同步的幂等键，使本地记录和云端同步使用同一份事件数据。")
    doc.add_heading("8.3 教师看板", level=2)
    add_body(doc, "看板按日期、题目、知识点和错误类型筛选，汇总尝试、提交、通过率、完成情况、常见错误和薄弱知识点。CSV导出使用UTF-8，导出文件应按学生数据管理要求保管。")
    page_break(doc)

    # 10 - Cloud
    add_page_title(doc, "9", "云端班级、任务与增量同步", "云端能力增强多教师协作，但本地核心教学不依赖网络可用性。")
    doc.add_heading("9.1 教师操作", level=2)
    decimal_id = create_numbering(doc, "decimal")
    for text in [
        "创建班级并选中班级，按邮箱添加教师或学生。",
        "填写本地题目ID和任务标题，创建班级任务。",
        "教师只管理自己参与的班级；管理员可查看全部班级。",
        "学生只能查看所属班级和分配给班级的任务，不能维护成员。",
    ]:
        add_list_item(doc, text, decimal_id)
    doc.add_heading("9.2 同步规则", level=2)
    add_body(doc, "学习事件以事件ID作为幂等键上传，服务端按版本号提供增量下载。重复同步同一事件不会产生重复记录。")
    add_body(doc, "桌面端只选择属于当前登录账号的事件。访客和其他账号的本地事件不会上传到当前账号；同一台机器上的事件按所有者字段隔离，同步时不做跨账号合并。网络失败时保留本地数据，用户可在网络恢复后再次同步。")
    add_callout(doc, "身份隔离", "退出登录后旧令牌不能继续访问服务；班级成员、任务和同步接口均在服务端重新鉴权。", PALE_AMBER)
    doc.add_heading("9.3 部署边界", level=2)
    add_body(doc, "正式教学数据应通过已完成ICP备案的域名和HTTPS访问云端服务。数据库端口和应用内部HTTP端口不应直接暴露给互联网。")
    page_break(doc)

    # 11 - Data
    add_page_title(doc, "10", "连接管理、备份与恢复", "数据库密码仅在当前连接调用中临时使用，运行数据与程序文件分开保存。")
    doc.add_heading("10.1 连接管理", level=2)
    add_body(doc, "设置页可维护SQLite、MySQL或MariaDB连接的非敏感配置并执行连接测试。密码只在当前连接测试会话中使用，不作为长期配置保存。访问真实课程数据库时应使用最小权限账号。")
    doc.add_heading("10.2 本地数据维护", level=2)
    decimal_id = create_numbering(doc, "decimal")
    for text in [
        "立即备份：保存应用数据库和演示数据库，并保留有限数量的最近快照。",
        "恢复备份：先校验SQLite完整性并备份当前状态，再恢复所选快照。",
        "恢复演示库：只重建演示数据，不影响题库、学习记录和连接设置。",
        "升级保护：数据库迁移按版本顺序执行，升级前创建可恢复副本。",
    ]:
        add_list_item(doc, text, decimal_id)
    add_callout(doc, "数据位置", "正式Windows版本默认使用本地应用数据目录；卸载程序默认保留用户数据。开发环境的app-data、数据库、日志和密钥文件不应复制到安装目录或对外共享。", PALE_BLUE)
    doc.add_heading("10.3 恢复后的检查", level=2)
    add_body(doc, "重新启动软件，确认首页可进入、演示表可查询、题目目录和学习记录存在；如完整性检查失败，停止覆盖并保留原备份。")
    page_break(doc)

    # 12 - Design originality
    add_page_title(doc, "11", "软件结构与关键机制", "本章说明软件的分层结构，以及各功能模块如何组合为面向SQL教学的完整系统。")
    doc.add_heading("11.1 分层结构", level=2)
    add_table(doc, ["层次", "职责", "代表实现"], [
        ("桌面层", "JavaFX页面、状态、异步交互和确认", "自由SQL、练习、AI、知识、云端控制器"),
        ("应用层", "用例契约、请求结果、学习事件", "执行、风险、练习、同步服务接口"),
        ("领域层", "题目规则和值对象", "难度、数据集、评测规则"),
        ("基础设施层", "SQLite/JDBC、AI、FTS5、备份", "风险分析器、查询执行器、评测引擎、FTS5检索模块、学情统计模块"),
        ("云端服务", "账号、角色、班级、任务、同步", "鉴权、成员关系、事件幂等接口"),
    ], [1600, 3780, 3980])
    doc.add_heading("11.2 关键机制组合", level=2)
    for text in [
        "统一安全路径：自由输入、AI草案和练习草稿都经过“拆分—风险分析—确认”流水线，任何来源都不能绕过安全规则。",
        "确定性教学评测：系统执行学生SQL和参考SQL，再按列、行、顺序、行数和关键结构比较结果；评测差异、提示和风险拦截写入同一套学习事件。",
        "受控知识检索：课程资料写入本地SQLite FTS5索引，只作为教学上下文使用，不能修改系统指令或放宽SQL安全规则。",
        "本地优先同步：核心数据先保存在本地，登录后仅同步当前所有者的事件；事件ID负责幂等去重，版本号用于增量下载。",
        "课堂连续性维护：恢复前校验SQLite完整性，迁移按版本执行，演示库可单独重建，避免影响题库和学习记录。",
    ]:
        item = add_list_item(doc, text, bullet_id)
        item.paragraph_format.space_after = Pt(2)
        item.paragraph_format.line_spacing = 1.0
        for run in item.runs:
            set_run_font(run, size=9.2, color=INK)
    doc.add_heading("11.3 机制对照", level=2)
    add_table(doc, ["维度", "简单实现方式", "本软件的处理方式"], [
        ("SQL执行控制", "按分号切分并匹配关键词", "状态机拆分、五类风险决策和高风险确认"),
        ("练习判定", "比较答案文本", "执行后按列、行、顺序、行数和结构规则评测"),
        ("AI使用", "模型生成后直接调用执行", "模型只生成结构化草案，Java侧复核后由用户决定"),
        ("数据同步", "全部数据集中上传", "本地优先，按所有者隔离并以事件ID幂等同步"),
    ], [1550, 3230, 4580])
    page_break(doc)

    # 13 - Troubleshooting and acceptance
    add_page_title(doc, "12", "故障处理与使用核对", "遇到故障时先确认本地数据和安全边界，再处理外部服务。")
    add_table(doc, ["现象", "处理方式"], [
        ("SQL语法错误", "检查表名、列名、引号、括号和逗号；先用简单SELECT缩小问题"),
        ("高风险语句被阻止", "阅读风险说明；禁止语句不可通过确认放行"),
        ("AI不可用", "检查Ollama或HTTPS服务配置；继续使用手写SQL"),
        ("课程资料导入失败", "确认UTF-8编码、文件类型、大小和读取权限"),
        ("云端同步失败", "保留本地记录，检查登录状态和网络后重试"),
        ("恢复失败", "停止覆盖，保留当前数据与原快照并检查SQLite完整性"),
    ], [2650, 6710])
    doc.add_heading("12.1 功能自检", level=2)
    for text in [
        "应用能够进入登录页；外部服务不可用时仍可使用访客模式。",
        "演示数据库可以执行只读查询，禁止语句会显示拦截原因。",
        "练习会话可以运行草稿、申请提示并提交确定性评测。",
        "AI不可用时页面给出错误提示，不影响手写SQL和本地练习。",
        "备份完成后保留生成时间和文件位置，恢复前先执行完整性检查。",
    ]:
        add_list_item(doc, text, bullet_id)
    page_break(doc)

    # 14 - Installation and initialization
    add_page_title(doc, "13", "安装、启动与首次初始化", "首次启动同时完成运行环境检查、本地目录准备和数据库结构初始化。")
    doc.add_heading("13.1 运行条件", level=2)
    add_table(doc, ["项目", "要求或说明"], [
        ("操作系统", "Windows 10或Windows 11 64位"),
        ("正式安装包", "包含自有Java运行时，普通用户无需单独安装JDK"),
        ("开发运行", "JDK 21、Maven 3.9及可用的JavaFX图形环境"),
        ("建议分辨率", "不低于1366×768；窄屏导航支持横向滚动"),
        ("可选服务", "本地Ollama、兼容OpenAI协议的HTTPS服务、学校云端服务"),
    ], [2350, 7010])
    doc.add_heading("13.2 首次启动过程", level=2)
    for text in [
        "应用检查JavaFX和运行环境，无法启动图形界面时返回明确的环境提示。",
        "创建应用数据目录，并把应用数据库与教学演示数据库分开保存。",
        "按版本顺序执行SQLite迁移，创建题目、练习、事件、知识索引和连接配置表。",
        "初始化演示数据库；若文件已存在，则保留用户数据并只执行必要升级。",
        "初始化成功后进入登录页；外部AI或云端不可用不会阻止本地入口启动。",
    ]:
        add_list_item(doc, text, decimal_id)
    add_callout(doc, "启动失败", "先检查数据目录写权限、磁盘空间和杀毒软件拦截。不要通过删除整个数据目录来处理单个数据库错误，应先保留备份。", PALE_AMBER)
    page_break(doc)

    # 15 - Schema and connection testing
    add_page_title(doc, "14", "表结构浏览与连接测试", "表结构既帮助学生理解数据，也为SQL草案提供受控、可核对的上下文。")
    doc.add_heading("14.1 浏览数据库结构", level=2)
    add_body(doc, "进入“表结构”后选择当前连接。页面读取数据库元数据，并按表展示字段名、类型、可空性和主键标记。读取失败时页面保留当前状态，同时显示便于用户理解的错误信息。")
    for text in [
        "先确认当前数据库与题目要求一致，避免在错误连接上执行练习。",
        "查看表名与字段拼写，确认主键、可空字段和数据类型。",
        "编写JOIN前核对关联字段；使用AI时只提供完成任务所需的结构。",
        "结构刷新属于只读操作，不会修改数据库对象。",
    ]:
        add_list_item(doc, text, bullet_id)
    doc.add_heading("14.2 测试外部连接", level=2)
    add_table(doc, ["连接类型", "需要填写", "安全建议"], [
        ("SQLite", "本地数据库文件路径", "确认文件属于当前课程，不使用未知数据库"),
        ("MySQL", "主机、端口、库名、用户名和临时密码", "使用专用只读账号，限制来源地址"),
        ("MariaDB", "主机、端口、库名、用户名和临时密码", "先测试连接，再执行只读查询"),
    ], [1700, 3760, 3900])
    add_callout(doc, "密码处理", "连接测试密码由调用方持有，只用于当前测试调用；服务结果、日志和长期连接配置均不得包含密码。", PALE_AMBER)
    page_break(doc)

    # 16 - Exercise authoring
    add_page_title(doc, "15", "教师题库与练习规则设计", "题目在本软件中是一个可执行的教学单元，由隔离数据集、练习目标、评测规则和提示层级四部分组成。")
    doc.add_heading("15.1 新建题目", level=2)
    for text in [
        "填写题目标题、知识点和难度，使学生能判断当前练习目标。",
        "选择隔离数据集并验证初始化SQL，确保每次会话得到一致起点。",
        "填写参考SQL用于教师核对；学生最终得分不直接依赖字符串相等。",
        "配置结果列、结果行、顺序、行数或关键结构等确定性判定规则。",
        "最多配置三层提示，从概念提醒逐步过渡到结构提示，避免第一层直接泄露答案。",
        "保存后先以教师身份试做，再启用题目供学生选择。",
    ]:
        add_list_item(doc, text, decimal_id)
    doc.add_heading("15.2 编辑与生命周期", level=2)
    add_table(doc, ["操作", "适用情形", "结果"], [
        ("编辑", "修正文案、规则或提示", "更新当前题目定义"),
        ("复制", "基于已有题目设计变式", "生成独立题目，保留原题"),
        ("停用", "暂不让学生开始新会话", "历史记录保留，目录不再开放"),
        ("重新启用", "规则复核完成", "题目重新进入可选目录"),
    ], [1700, 3700, 3960])
    add_callout(doc, "设计建议", "一题集中训练一个主要知识点。判定规则应允许语义等价SQL，不应把空格、大小写或无关别名差异误判为错误。", PALE_GREEN)
    page_break(doc)

    # 17 - Package import/export
    add_page_title(doc, "16", "题包导入、导出与冲突处理", "题包用于课程迁移和教师协作，导入过程以整体一致性为先。")
    doc.add_heading("16.1 导出题包", level=2)
    add_body(doc, "教师从题库选择需要共享的题目后导出版本化JSON题包。题包包含题目定义、数据集引用、参考SQL、评测规则和提示，不包含学生作答、连接密码、AI密钥或运行日志。")
    doc.add_heading("16.2 导入题包", level=2)
    for text in [
        "选择来源明确的题包文件，导入前保留当前题库备份。",
        "系统解析版本和结构，校验题目ID、必填字段、难度、规则和提示数量。",
        "对文件内重复ID、与本地题目冲突或无效数据集引用给出明确错误。",
        "全部题目验证通过后才写入数据库；任一题失败时整批回滚。",
        "导入完成后检查新增、更新和跳过数量，并抽查关键题目的评测结果。",
    ]:
        add_list_item(doc, text, decimal_id)
    doc.add_heading("16.3 版本与可追溯性", level=2)
    add_body(doc, "题包格式带有版本信息，使后续字段扩展能够被识别。教师应保留导出时间、课程名称和修改说明，不要用聊天软件中的未知附件直接覆盖正式题库。")
    add_callout(doc, "事务边界", "导入采用数据库事务控制，目标是避免“部分题目已写入、部分题目失败”的半完成状态。", PALE_BLUE)
    page_break(doc)

    # 18 - Analytics and export
    add_page_title(doc, "17", "教师看板、筛选与CSV导出", "看板在总次数之外，还提供按题目、知识点和错误类型的细分统计，便于定位具体教学问题。")
    doc.add_heading("17.1 查看学情", level=2)
    for text in [
        "先选择日期范围，避免把不同教学阶段混入同一统计。",
        "按题目或知识点定位通过率低、重复尝试多的内容。",
        "按错误类型区分语法问题、风险拦截、结果差异和连接故障。",
        "结合提示使用次数判断学生是在概念、结构还是细节层面遇到困难。",
        "打开具体记录时只查看完成教学判断所需的数据，不传播学生身份信息。",
    ]:
        add_list_item(doc, text, decimal_id)
    doc.add_heading("17.2 指标解释", level=2)
    add_table(doc, ["指标", "含义", "使用提醒"], [
        ("尝试数", "运行草稿或开始作答的次数", "不能单独代表学习质量"),
        ("提交数", "进入正式评测的次数", "与尝试数结合观察"),
        ("通过率", "通过提交占有效提交比例", "筛选范围不同不可直接比较"),
        ("薄弱知识点", "失败、提示和重复尝试集中项", "用于安排复习或示范"),
    ], [1900, 3500, 3960])
    add_callout(doc, "CSV导出", "导出文件采用UTF-8编码。文件离开应用后由用户负责访问控制、存储期限和脱敏，不应上传到公开网盘。", PALE_AMBER)
    page_break(doc)

    # 19 - Classroom workflow
    add_page_title(doc, "18", "一次完整课堂的使用流程", "本章按课前、课中、课后三个阶段，给出一节完整课的操作顺序。")
    doc.add_heading("18.1 课前准备", level=2)
    for text in [
        "启动应用并检查演示数据库可查询，创建一份手动备份。",
        "复核本节课题目、数据集、评测规则与提示层级。",
        "如需AI，提前检查本地模型；如需云端，确认已完成ICP备案的域名和HTTPS可用。",
        "使用教师测试账号走通登录、任务发布和学生查看路径。",
    ]:
        add_list_item(doc, text, bullet_id)
    doc.add_heading("18.2 课堂进行", level=2)
    add_body(doc, "教师先用表结构页说明数据，再让学生在自由SQL中完成低风险查询。进入闯关练习后，学生运行草稿、使用提示并提交；教师通过看板观察共性错误，对高频问题进行现场讲解。AI助手只用于展示草案形成和复核过程，不代替学生提交。")
    doc.add_heading("18.3 课后处理", level=2)
    for text in [
        "导出必要的统计结果并按课程数据要求保存。",
        "整理高频错误和薄弱知识点，调整下一次题目或提示。",
        "同步已登录学生的学习事件；网络失败时保留本地记录稍后重试。",
        "保留题包与备份，清理临时导出文件和脱敏截图。",
    ]:
        add_list_item(doc, text, bullet_id)
    add_callout(doc, "课堂连续性", "AI、外部数据库或云端暂时不可用时，演示库、手写SQL、隔离练习和本地记录仍能支持核心课堂活动。", PALE_GREEN)
    page_break(doc)

    # 20 - Security boundaries
    add_page_title(doc, "19", "安全与数据边界", "本章汇总AI、SQL执行、连接信息、课程资料和学习数据的使用范围。")
    doc.add_heading("19.1 对象与限制", level=2)
    add_table(doc, ["对象", "允许范围", "禁止或限制"], [
        ("AI模型", "生成结构化草案和解释", "不得执行SQL、持有JDBC连接或改变安全规则"),
        ("SQL执行", "单条、通过风险检查的语句", "禁止语句直接拦截；高风险操作须确认"),
        ("连接密码", "当前测试调用内临时使用", "不得写入配置、日志、截图或导出文件"),
        ("课程资料", "本地检索和受控上下文", "不能作为系统指令或放宽SQL安全"),
        ("学习数据", "本地记录和本人账号同步", "访客及其他账号事件不得混传"),
    ], [1750, 3840, 3770])
    add_callout(doc, "版本说明", f"本手册描述的软件功能以{VERSION}版本为准。", PALE_BLUE)

    props = doc.core_properties
    props.title = f"{SOFTWARE_NAME}用户手册{VERSION}"
    props.subject = "SQLTeacher V1.1.0用户操作手册"
    props.author = "河南科技大学"
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(MANUAL_PATH)
    return MANUAL_PATH


def git_text(path):
    result = subprocess.run(
        ["git", "show", f"{BASELINE}:{path}"],
        cwd=ROOT,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        check=True,
    )
    return result.stdout.decode("utf-8")


def build_representative_source():
    compiled = []
    for path in REPRESENTATIVE_FILES:
        compiled.append(f"// FILE: {path}")
        compiled.extend(git_text(path).splitlines())
    if len(compiled) < 3000:
        raise RuntimeError(f"Representative source has only {len(compiled)} lines")
    selected = compiled[:1500] + compiled[-1500:]

    doc = Document()
    configure_a4(doc.sections[0], top=1.35, bottom=1.25, left=1.25, right=1.15)
    style = doc.styles["Normal"]
    style.font.name = "Courier New"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(5.8)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    style.paragraph_format.line_spacing = Pt(10.1)

    for page in range(60):
        if page:
            doc.add_page_break()
        for line in selected[page * 50:(page + 1) * 50]:
            p = doc.add_paragraph()
            p.paragraph_format.keep_together = True
            p.add_run(line)

    props = doc.core_properties
    props.title = f"{SOFTWARE_NAME}源程序鉴别材料{VERSION}"
    props.subject = "SQL安全、练习评测、AI安全、知识检索、学情分析与备份恢复代表性源程序"
    props.author = "河南科技大学"
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(SOURCE_PATH)
    return SOURCE_PATH, len(compiled), selected


def main():
    manual = build_manual()
    source, compiled_count, selected = build_representative_source()
    print(f"Created: {manual}")
    print(f"Created: {source}")
    print(f"Baseline: {BASELINE}")
    print(f"Representative files: {len(REPRESENTATIVE_FILES)}")
    print(f"Compiled lines: {compiled_count}; submitted lines: {len(selected)}")


if __name__ == "__main__":
    main()
